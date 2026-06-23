import docx
import sys

def main():
    doc_path = r"c:\Users\USUARIO\Documents\Aplicaciones\OnProduction\documento_requerimiento\requerimientos.docx"
    doc = docx.Document(doc_path)
    
    print("=== TEXTO DEL DOCUMENTO ===")
    for para in doc.paragraphs:
        if para.text.strip():
            print(para.text)
            
    print("\n=== TABLAS DEL DOCUMENTO ===")
    for i, table in enumerate(doc.tables):
        print(f"\nTablas #{i+1}:")
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            print(" | ".join(row_data))

if __name__ == "__main__":
    main()
