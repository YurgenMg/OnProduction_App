import docx

def debug_docx_runs():
    doc_path = r"c:\Users\USUARIO\Documents\Aplicaciones\OnProduction\documento_requerimiento\requerimientos.docx"
    doc = docx.Document(doc_path)
    
    print("=== PÁRRAFOS ===")
    for i, p in enumerate(doc.paragraphs):
        print(f"P{i}: text='{p.text}' style='{p.style.name}'")
        for r in p.runs:
            print(f"  Run: text='{r.text}' bold={r.bold} italic={r.italic}")
            
    print("\n=== TABLAS ===")
    for i, t in enumerate(doc.tables):
        print(f"Tabla {i}: rows={len(t.rows)}, cols={len(t.columns)}")
        for r_idx, row in enumerate(t.rows):
            cells = [c.text.strip().replace('\n', ' ') for c in row.cells]
            print(f"  Fila {r_idx}: {cells}")

if __name__ == "__main__":
    debug_docx_runs()
